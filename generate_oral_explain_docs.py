from docx import Document
from docx.shared import Pt


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_function_block(doc: Document, item: dict) -> None:
    doc.add_heading(item["title"], level=2)

    doc.add_paragraph("A. Giải thích theo Sequence (dễ trả lời khi thầy hỏi luồng API)")
    for line in item["sequence"]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph("B. Giải thích theo Activity (dễ trả lời khi thầy hỏi điều kiện/rẽ nhánh)")
    for line in item["activity"]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph("")


def build_app_doc(path: str) -> None:
    doc = Document()
    set_default_font(doc)
    doc.add_heading("Giải thích Sequence + Activity - Chức năng App SmartTour", level=1)
    doc.add_paragraph(
        "Tài liệu học nhanh để vấn đáp: mỗi chức năng tách 2 phần Sequence và Activity, "
        "kèm hàm + file chạy thực tế trong code."
    )

    functions = [
        {
            "title": "1) Quét QR để vào app (QR Gate)",
            "sequence": [
                "Nhận diện QR diễn ra ở hàm `OnBarcodesDetected(...)` trong file `SmartTourApp/Pages/QrGatePage.xaml.cs`.",
                "Kiểm tra hợp lệ QR dùng `IsValidGateQr(...)`, chuẩn hóa deep link dùng `NormalizeQr(...)` trong cùng file.",
                "Lưu phiên 7 ngày bằng `Preferences.Default.Set(...)` tại `OnBarcodesDetected(...)` (key `qr_gate_until_utc`).",
                "Điều hướng sau quét thành công: tạo `LoadingPage` hoặc `AppShell` ngay trong `OnBarcodesDetected(...)`.",
                "Sau điều hướng, app gửi presence bằng `ApiService.PostPresenceHeartbeatAsync()` trong `SmartTourApp/Services/ApiService.cs`.",
            ],
            "activity": [
                "Nhánh xin quyền camera: `EnsureCameraPermissionAsync()` quyết định cho phép hay dừng scan.",
                "Nhánh QR sai/đúng: trong `OnBarcodesDetected(...)`, QR sai thì bật scan lại; QR đúng thì khóa camera và đi tiếp.",
                "Nhánh điều hướng: nếu resolve được service (`PoiRepository`, `TrackingService`) thì vào `LoadingPage`; nếu không thì vào `AppShell`.",
                "Tránh xử lý trùng bằng biến `_processing` + `_navigationStarted` trong `QrGatePage`.",
            ],
        },
        {
            "title": "2) Khởi động app + tải POI + bật tracking",
            "sequence": [
                "Luồng chính nằm trong `InitAsync()` của `SmartTourApp/Pages/LoadingPage.xaml.cs`.",
                "Tải POI gọi `repo.GetPois()`; hàm thật là `PoiRepository.GetPois()` trong `SmartTourApp/Services/PoiRepository.cs`.",
                "Tracking bắt đầu bằng `tracking.Start()` trong `LoadingPage.InitAsync()`.",
                "Presence heartbeat gửi ngay sau boot qua `api.PostPresenceHeartbeatAsync()` trong `InitAsync()`.",
                "Kết thúc luồng bằng chuyển `Application.Current.MainPage = new AppShell()` trong `InitAsync()`.",
            ],
            "activity": [
                "Nếu API không có dữ liệu, `PoiRepository.GetPois()` rẽ nhánh fallback SQLite local (`db.GetPois()`).",
                "Trong `InitAsync()`, app vẫn cho qua dù POI rỗng để tránh crash khi offline.",
                "Auto-play đầu phiên có điều kiện theo `SettingsPage.AutoPlayKey` (bật thì chạy, tắt thì bỏ qua).",
            ],
        },
        {
            "title": "3) Lấy danh sách POI theo ngôn ngữ (`GET /api/pois?lang=xx`)",
            "sequence": [
                "App gọi API tại `ApiService.GetPois()` trong `SmartTourApp/Services/ApiService.cs`.",
                "URL thực tế tạo ở `GetPois()` là `api/pois?lang=...` (đúng như sơ đồ sequence).",
                "Backend nhận ở `PoisController.GetPois(...)` trong `SmartTourBackend/Controllers/PoisController.cs`.",
                "Bước `Query Poi + Food + translation` nằm trong `GetPois(...)`: `.Include(p => p.Foods).ThenInclude(f => f.FoodTranslations)...`.",
                "Backend trả về POI đã lọc active/approved; app nhận về và tiếp tục cache trong `PoiRepository.GetPois()`.",
            ],
            "activity": [
                "Activity nhánh tìm kiếm/lọc xảy ra trong `PoisController.GetPois(...)` với query `q`, `lat/lng/maxDistanceKm`.",
                "Nếu online: `PoiRepository.GetPois()` dùng server data + ghi xuống SQLite.",
                "Nếu offline/API fail: `PoiRepository.GetPois()` trả dữ liệu local cache.",
            ],
        },
        {
            "title": "4) Xem POI map/list + chọn POI",
            "sequence": [
                "Map khởi tạo tại `MapPage.InitMap()` trong `SmartTourApp/Pages/MapPage.xaml.cs`.",
                "Nạp dữ liệu POI bằng `pois = await repo.GetPois()` trong `InitMap()`.",
                "Render POI layer qua `vm.LoadPois(...)` và thêm layer vào map trong `InitMap()`.",
                "Khi tap marker, sự kiện `OnMapInfo(...)` chọn POI gần nhất và gọi `SelectPoi(...)`.",
                "Mở chi tiết POI bằng `PoiCard_Tapped(...)` -> `Shell.Current.GoToAsync(nameof(PoiDetailPage), ...)`.",
            ],
            "activity": [
                "Nếu mở app từ deep link có `TargetPoiId`, `OnAppearing()` rẽ nhánh tự zoom đến POI đích.",
                "Nếu user tự đóng card (`cardManuallyClosed`), activity không tự bật lại ngay cho đến vòng cập nhật phù hợp.",
                "Nếu mất GPS nhưng có offline center thì map vẫn center từ cache (`TryGetOfflineCenter(...)`).",
            ],
        },
        {
            "title": "5) Nghe audio theo ngôn ngữ",
            "sequence": [
                "Lấy audio/translation từ API qua `ApiService.GetPoiAudios(...)` và `ApiService.GetTtsScripts(...)`.",
                "Endpoint map đúng sequence: `GET /api/audio/poi/{poiId}` ở `SmartTourBackend/Controllers/AudioController.cs` hàm `GetPoiAudios(...)`.",
                "Trong map card, tên đa ngôn ngữ có thể làm mới qua `MapPage.RefreshPoiCardNameAsync(...)` gọi `api.GetTtsScripts(poi.Id)`.",
                "Khi phát xong hoặc nghe đủ ngưỡng, app gửi sự kiện listen qua `ApiService.PostPoiAudioListenAsync(...)`.",
            ],
            "activity": [
                "Nếu không có audio cloud, app fallback local/cache theo luồng trong repository + offline sync.",
                "Backend kiểm tra ngưỡng nghe ở `AnalyticsController.IsQualifiedListen(...)` rồi mới nhận listen hợp lệ.",
                "Nhánh duplicate 15s bị loại tại `AudioListenIngestionService.TryEnqueue(...)`.",
            ],
        },
        {
            "title": "6) Offline map/audio + đồng bộ lại",
            "sequence": [
                "Tải map offline ở `MapPage.StartDownloadAsync(...)` gọi `OfflineMapService.DownloadAreaForTourAsync(...)`.",
                "Seed nhanh tile khi mở map tại `MapPage.EnsureSeedTilesAsync()`.",
                "Prefetch dữ liệu offline (scripts/foods/audio) chạy trong `PoiRepository.GetPois()` bằng `offlineSync.PrefetchPoiDataAsync(...)`.",
                "Khi online lại, queue/offline data được xử lý trong `OfflineSyncService` (file `SmartTourApp/Services/OfflineSyncService*.cs`).",
            ],
            "activity": [
                "Nếu không mạng, `OnDownloadMapTapped(...)` rẽ nhánh báo lỗi và không bắt đầu download.",
                "Nếu map đã cache đủ thì rẽ nhánh thông báo 'đã tải' và dừng.",
                "Ở `PoiRepository.GetPois()`, online thì sync server trước; fail thì fallback SQLite.",
            ],
        },
        {
            "title": "7) Gửi playlog và route",
            "sequence": [
                "Playlog gửi qua `ApiService.PostPlayLog(...)` -> endpoint backend `POST /api/pois/playlog` (`PoisController.PostPlayLog(...)`).",
                "Route session gửi qua `ApiService.PostRouteSession(...)` -> `POST /api/routes/session` (`RouteSessionController.PostSession(...)`).",
                "State machine route nằm ở `RouteTrackingService` (`OnLocationUpdatedAsync`, `OnManualAudioPlayedAsync`, `FlushSessionAsync`).",
                "Khi app đóng, flush route bằng `RouteTrackingService.FlushOnAppClosingAsync()`.",
            ],
            "activity": [
                "Route chỉ ghi khi đủ điều kiện dwell/audio-manual theo rule trong `RouteTrackingService`.",
                "Nếu session timeout 90 phút, `RecoverSessionOnStartupAsync()`/`OnLocationUpdatedAsync()` sẽ flush rồi reset.",
                "Nếu stop < ngưỡng tối thiểu, `FlushSessionAsync()` sẽ bỏ session để tránh nhiễu.",
            ],
        },
        {
            "title": "8) Heartbeat/offline thiết bị",
            "sequence": [
                "App gửi heartbeat bằng `ApiService.PostPresenceHeartbeatAsync()`.",
                "App gửi offline bằng `ApiService.PostPresenceOfflineAsync()`.",
                "Backend nhận ở `PresenceController.Heartbeat(...)` và `PresenceController.Offline(...)` trong `SmartTourBackend/Controllers/PresenceController.cs`.",
                "CMS dashboard đọc trạng thái từ `HomeController.GetDeviceStatus()` (`SmartTourCMS/Controllers/HomeController.cs`).",
            ],
            "activity": [
                "Loading và QrGate đều có nhánh gửi heartbeat khi vào app thành công.",
                "Nếu lỗi mạng khi gửi presence, `ApiService` bắt exception và bỏ qua để không crash UI.",
                "Dashboard tự refresh theo chu kỳ ở view Home nên trạng thái online/offline cập nhật liên tục.",
            ],
        },
        {
            "title": "9) Đổi ngôn ngữ app",
            "sequence": [
                "Luồng nằm ở `SettingsPage.Save(...)` trong `SmartTourApp/Pages/SettingsPage.xaml.cs`.",
                "Gán ngôn ngữ mới qua `lang.Current = selected` (LanguageService).",
                "Sau đó reset cache và engine: `repo.ClearCache()`, `tracking.Stop()`, `narration.Reset()`, `geo.Reset()`, `routeTracking.Reset()`.",
                "Cuối luồng reload app bằng `Application.Current.MainPage = new LoadingPage(repo, tracking)`.",
            ],
            "activity": [
                "Nếu ngôn ngữ không đổi, rẽ nhánh báo `LanguageUnchanged` và dừng.",
                "Nếu offline, hiển thị thông báo riêng cho chế độ cache local (`GetOfflineLangChangeMessage`).",
                "Trước khi reset route, `Save(...)` gọi `routeTracking.FlushOnAppClosingAsync()` để không mất session.",
            ],
        },
        {
            "title": "10) Xóa phiên QR trong Settings",
            "sequence": [
                "Nút xóa gọi `OnClearQrSessionTapped(...)` trong `SettingsPage.xaml.cs`.",
                "Xóa key session bằng `Preferences.Default.Remove(\"qr_gate_until_utc\")`.",
                "Lần mở app sau sẽ đi lại luồng `QrGatePage` vì không còn phiên hợp lệ.",
            ],
            "activity": [
                "Có nhánh xác nhận người dùng trước khi xóa (`DisplayAlert` confirm).",
                "Nếu user bấm Cancel thì không xóa gì.",
            ],
        },
        {
            "title": "11) Device token + bảo vệ API app",
            "sequence": [
                "Mọi call app đều qua `ApiService.EnsureDeviceTokenAsync()` trước khi gọi endpoint chính.",
                "Hàm này gọi `POST /api/auth/device-token` tại backend.",
                "Backend xử lý ở `AuthController.DeviceToken(...)` + cache token tại `DeviceTokenCacheService.GetOrCreateAsync(...)`.",
                "JWT được tạo bằng `JwtTokenService.CreateTokenAsync(...)` và gắn vào `Authorization` header trong ApiService.",
            ],
            "activity": [
                "Nếu token còn hạn (`_tokenExpiryUtc`) thì bỏ qua xin lại token.",
                "Nếu token sắp hết hạn, tự xin token mới rồi mới gọi API nghiệp vụ.",
            ],
        },
        {
            "title": "12) Ingestion lượt nghe audio qua queue",
            "sequence": [
                "App gửi `POST /api/analytics/poi-audio-listen` từ `ApiService.PostPoiAudioListenAsync(...)`.",
                "Backend nhận ở `AnalyticsController.PostPoiAudioListen(...)`.",
                "Controller đẩy queue bằng `IAudioListenIngestionService.TryEnqueue(...)`.",
                "Worker nền `AudioListenIngestionWorker.ExecuteAsync()` đọc batch và ghi DB bằng `FlushBatchAsync(...)`.",
            ],
            "activity": [
                "Nếu không đạt ngưỡng nghe, `PostPoiAudioListen(...)` trả `accepted=false` ngay.",
                "Nếu trùng trong 15 giây, `TryEnqueue(...)` trả reason `duplicate_window_15s`.",
                "Nếu queue full, trả reason `server_busy_queue_full`.",
            ],
        },
    ]

    for item in functions:
        add_function_block(doc, item)

    doc.save(path)


def build_web_doc(path: str) -> None:
    doc = Document()
    set_default_font(doc)
    doc.add_heading("Giải thích Sequence + Activity - Chức năng Web/CMS SmartTour", level=1)
    doc.add_paragraph(
        "Tài liệu học nhanh vấn đáp cho phần Web/CMS + Backend API. "
        "Mỗi chức năng có phần Sequence và Activity, chỉ rõ hàm/file."
    )

    functions = [
        {
            "title": "1) Đăng nhập CMS và phân quyền",
            "sequence": [
                "Xử lý login ở `AccountController.Login(...)` trong `SmartTourCMS/Controllers/AccountController.cs`.",
                "Hàm gọi `SignInManager.PasswordSignInAsync(...)`, role lấy từ Identity.",
                "Seed role Admin/Vendor được đảm bảo khi khởi động trong `SmartTourCMS/Program.cs`.",
                "Sau login thành công, user vào dashboard theo quyền (Admin/Vendor).",
            ],
            "activity": [
                "Sai tài khoản/mật khẩu thì rẽ nhánh trả lại form login.",
                "Đăng nhập đúng thì tạo cookie auth và điều hướng trang chính.",
                "Đổi mật khẩu có luồng riêng tại `AccountController.ChangePassword(...)`.",
            ],
        },
        {
            "title": "2) Tạo POI",
            "sequence": [
                "Màn tạo POI xử lý ở `PoiController.Create(Poi, IFormFile, List<IFormFile>)` trong `SmartTourCMS/Controllers/PoiController.cs`.",
                "Upload ảnh qua Cloudinary trong cùng hàm `Create(...)`.",
                "POI lưu DB bằng `_context.Pois.Add(poi)` rồi `SaveChangesAsync()`.",
                "Nếu là Admin thì duyệt luôn; Vendor thì set `ApprovalStatus = pending`.",
            ],
            "activity": [
                "Nhánh Admin: POI active ngay, đi luồng thành công.",
                "Nhánh Vendor: POI chờ duyệt, chỉ thông báo 'đã gửi duyệt'.",
                "Nếu thiếu mô tả, hàm `GenerateScriptWithAI(...)` có thể sinh script hỗ trợ.",
            ],
        },
        {
            "title": "3) Sửa POI",
            "sequence": [
                "GET edit ở `PoiController.Edit(int? id)`, POST edit ở `PoiController.Edit(int id, Poi...)`.",
                "Nếu Admin sửa trực tiếp, dữ liệu update ngay và có thể rebuild translation/audio khi đổi mô tả.",
                "Nếu Vendor sửa, hệ thống tạo yêu cầu pending bằng `ApprovalNote` để Admin duyệt.",
            ],
            "activity": [
                "Nhánh Vendor: chuyển về pending, không cập nhật live ngay.",
                "Nhánh Admin: giữ trạng thái approved và có thể regenerate translation/audio.",
                "Nếu không đúng quyền sở hữu POI thì trả `Forbid()`.",
            ],
        },
        {
            "title": "4) Xóa POI",
            "sequence": [
                "Xử lý ở `PoiController.Delete(int id)`.",
                "Hàm xóa dữ liệu liên quan: `PoiTranslations`, `PoiImages`, `Food`, `FoodTranslations`, `PlayLog`, `HeatmapEntries`.",
                "Cuối cùng xóa POI rồi `SaveChangesAsync()`.",
            ],
            "activity": [
                "Nếu POI không tồn tại -> `NotFound()`.",
                "Nếu không đúng quyền -> `Forbid()`.",
                "Nếu có exception, nhánh catch ghi lỗi vào `TempData[\"Error\"]`.",
            ],
        },
        {
            "title": "5) Quản lý translation POI",
            "sequence": [
                "Danh sách translation ở `TranslationController.Index(int poiId)`.",
                "Sửa translation ở `TranslationController.Edit(...)` (GET/POST).",
                "Chi tiết translation ở `TranslationController.Details(int id)`.",
            ],
            "activity": [
                "Người dùng vào trang translation -> tải danh sách theo `poiId`.",
                "Chỉnh sửa -> validate -> lưu DB -> quay lại danh sách.",
                "Xóa translation có nhánh xác định lại `poiId` để điều hướng đúng.",
            ],
        },
        {
            "title": "6) Generate/Regenerate audio",
            "sequence": [
                "Luồng CMS gọi `_voiceService.GenerateAndUploadAudio(...)` ở `PoiController.RegenerateTranslationAudios(...)`.",
                "Interface `IVoiceService` nằm ở `SmartTourBackend/Service/IVoiceService.cs`.",
                "Implement Azure TTS + Cloudinary nằm ở `SmartTourBackend/Service/VoiceService.cs` hàm `GenerateAndUploadAudio(...)`.",
                "Backend API cũng có endpoint regenerate: `AudioController.RegeneratePoiAudios(...)` và `GenerateForTranslation(...)`.",
            ],
            "activity": [
                "Nếu translation đã có audio thì bỏ qua (skip).",
                "Nếu generate lỗi, đếm `missing` và báo warning cho user.",
                "Nếu cấu hình Azure/Cloudinary thiếu, service trả rỗng và ghi log lỗi.",
            ],
        },
        {
            "title": "7) Dashboard POI stats (lượt nghe)",
            "sequence": [
                "CMS gọi API trong `HomeController.GetDashboardPoiStats()`.",
                "Query thực hiện trên `PlayLog` + `Poi` và group theo POI trong chính hàm này.",
                "Backend API thống kê chuyên sâu tương ứng là `PoisController.GetListenStats()`.",
            ],
            "activity": [
                "Nhánh Admin: xem toàn bộ dữ liệu.",
                "Nhánh Vendor: lọc dữ liệu theo POI thuộc vendor.",
                "Kết quả trả về để vẽ chart trên dashboard.",
            ],
        },
        {
            "title": "8) Heatmap và route phổ biến",
            "sequence": [
                "CMS mở heatmap qua `HeatmapController.GetHeatmapData()` (`SmartTourCMS/Controllers/HeatmapController.cs.cs`).",
                "Backend nhận `GET /api/heatmap` tại `HeatmapController.GetHeatmap()`.",
                "Route phổ biến lấy ở `RouteSessionController.GetPopularRoutes(...)`.",
            ],
            "activity": [
                "Nếu không có dữ liệu, API trả danh sách rỗng nhưng vẫn success.",
                "Có dữ liệu thì group/sort rồi trả về cho map chart.",
            ],
        },
        {
            "title": "9) Quản lý user hệ thống",
            "sequence": [
                "Danh sách user ở `UserController.Index(...)`.",
                "Tạo vendor ở `UserController.CreateVendor(...)`.",
                "Khóa/mở user ở `UserController.ToggleLock(...)`.",
                "Reset mật khẩu ở `UserController.ResetPassword(...)`.",
            ],
            "activity": [
                "Admin thao tác user -> Identity update -> trả thông báo thành công/thất bại.",
                "Nếu id không hợp lệ hoặc user không tồn tại thì rẽ nhánh lỗi phù hợp.",
            ],
        },
        {
            "title": "10) Dashboard thiết bị online",
            "sequence": [
                "Trang dashboard lấy số online ở `HomeController.Index()` bằng `devices.Count(x => x.IsActive)`.",
                "API refresh định kỳ nằm ở `HomeController.GetDeviceStatus()`.",
                "Dữ liệu thiết bị lấy từ `GetDeviceStatusesSafeAsync(...)` query `_context.DevicePresences`.",
                "App gửi heartbeat/offline tới backend ở `PresenceController.Heartbeat(...)` và `Offline(...)`.",
            ],
            "activity": [
                "Online được tính theo ngưỡng `ActiveThresholdSeconds` (20 giây).",
                "Nếu bảng `DevicePresences` chưa tồn tại, nhánh catch trả list rỗng và log warning.",
                "Vendor không phải Admin bị chặn xem trạng thái thiết bị (`Forbid()`).",
            ],
        },
        {
            "title": "11) Tạo thanh toán Premium MoMo (queue)",
            "sequence": [
                "CMS bấm tạo payment ở `PremiumController.CreatePayment(...)`.",
                "CMS gọi backend proxy endpoint `POST /api/vendor/premium/create-payment-cms`.",
                "Backend xử lý tại `VendorPremiumController.CreatePaymentFromCms(...)` -> `CreatePaymentCore(...)`.",
                "Order được enqueue bằng `_paymentQueue.TryEnqueue(...)` (service `MoMoPaymentQueueService`).",
                "Worker nền `MoMoPaymentWorker.ProcessOneAsync(...)` gọi MoMo create và cập nhật status DB.",
            ],
            "activity": [
                "Nếu queue full -> nhánh failed `queue_full`, trả 429.",
                "Nếu MoMo trả URL hợp lệ -> status `awaiting_payment`.",
                "Nếu MoMo lỗi hoặc HTTP lỗi -> status `failed` + `LastError`.",
            ],
        },
        {
            "title": "12) Nhận IPN MoMo + kích hoạt Premium",
            "sequence": [
                "MoMo callback vào `VendorPremiumController.HandleMomoIpn(...)`.",
                "Hàm verify chữ ký bằng `BuildIpnSignatureString(...)` + `Sign(...)`.",
                "Nếu hợp lệ và `resultCode = 0`, gọi `ApplyPremiumBenefitAsync(...)` để tăng hạn Premium cho POI.",
                "CMS trang return đọc trạng thái qua `PremiumController.PaymentReturn(...)`.",
            ],
            "activity": [
                "Sai chữ ký -> status order `failed` với `invalid_signature`.",
                "IPN success -> order `paid`, set `PaidAt`, cập nhật `Poi.IsPremium` + `PremiumExpiresAt`.",
                "IPN fail -> giữ failed và lưu `LastError` để tra soát.",
            ],
        },
        {
            "title": "13) Poll trạng thái thanh toán Premium",
            "sequence": [
                "CMS AJAX gọi `PremiumController.GetPaymentStatus(orderId, forceProviderCheck)`.",
                "CMS gọi backend endpoint `POST /api/vendor/premium/order-status-cms`.",
                "Backend xử lý ở `VendorPremiumController.GetOrderStatusFromCms(...)` -> `GetOrderStatusCore(...)`.",
                "Nếu cần đối soát nhà cung cấp, `GetOrderStatusCore(...)` gọi `SyncOrderFromMoMoAsync(...)`.",
            ],
            "activity": [
                "Nếu order chưa paid và `forceProviderCheck = true`, activity rẽ nhánh query MoMo.",
                "Nếu query thành công -> đồng bộ trạng thái paid/failed về DB.",
                "Nếu query lỗi -> giữ trạng thái cũ và cập nhật `LastError`.",
            ],
        },
        {
            "title": "14) Ghi nhận listen analytics qua queue",
            "sequence": [
                "Endpoint nhận sự kiện ở `AnalyticsController.PostPoiAudioListen(...)`.",
                "Controller gọi `_ingestion.TryEnqueue(...)` từ `AudioListenIngestionService`.",
                "Worker nền `AudioListenIngestionWorker.ExecuteAsync()` gom batch và `FlushBatchAsync(...)` vào DB.",
                "Thống kê đọc ở `AnalyticsController.GetPoiAudioListenStats(...)`.",
            ],
            "activity": [
                "Qua ngưỡng nghe mới được nhận (`IsQualifiedListen(...)`).",
                "Duplicate trong 15s bị reject (reason `duplicate_window_15s`).",
                "Queue đầy thì reject để bảo vệ hệ thống khi tải cao.",
            ],
        },
    ]

    for item in functions:
        add_function_block(doc, item)

    doc.save(path)


def main() -> None:
    build_app_doc("GiaiThich_Sequence_Activity_App_SmartTour.docx")
    build_web_doc("GiaiThich_Sequence_Activity_Web_SmartTour.docx")
    print("Generated: GiaiThich_Sequence_Activity_App_SmartTour.docx")
    print("Generated: GiaiThich_Sequence_Activity_Web_SmartTour.docx")


if __name__ == "__main__":
    main()
